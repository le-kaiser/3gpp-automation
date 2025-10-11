import requests
from bs4 import BeautifulSoup
import pandas as pd
import zipfile
import os
from urllib.parse import urljoin
import docx
from datetime import datetime
import re

# --- Configuration ---
BASE_URL = "https://www.3gpp.org/ftp/tsg_ran/TSG_RAN"
SPEC_NUMBER = "38.101-1"
# Using a set for efficient lookups
CLAUSES_DATABASE = {'4.3', '5.1', '5.2', '5.3.1', '5.3.2', '5.3.3', '5.3.5', '6.3.2', '6.3.3', '6.3.3.1', '6.3.3.2', '6.5.1', '6.5.2.2', '6.5.2.1', '6.5.2.3', '6.5.2.4','6.5.2.3.1', '6.5.2.3.2', '6.5.2.3.3', '6.5.2.3.4', '6.5.2.3.7', '6.5.2.3.8', '6.5.2.3.9', '6.4' , '6.4.1', '6.4.2', '6.4.2.0', '6.4.2.1', '6.4.2.1a', '6.4.2.2', '6.4.2.3', '6.4.2.4', '6.4.2.4.1', '6.4.2.4.2','6.4.2.5', 'A.3','C.2','F.0','F.1','F.2','F.3','F.4','F.5','F.5.1','F.5.2','F.5.3','F.5.4','F.5.5','F.6','F.7','F.8','F.9', 'F.10', '6.5.1', '6.5.2.4', '5.3.6', 'D.2'} 
OUTPUT_FILE = "approved_clauses.xlsx"
TEMP_DIR = "temp_files"

# --- Main Functions ---

def get_sorted_meeting_folders(url):
    """
    Fetches and sorts the TSG-RAN meeting folders from the 3GPP website by date.
    """
    print(f"Fetching and sorting meeting folders by date from: {url}")
    try:
        response = requests.get(url)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching URL: {e}")
        return []

    soup = BeautifulSoup(response.text, 'html.parser')
    folders_with_dates = []

    # Find all table rows in the body
    for row in soup.find('tbody').find_all('tr'):
        # Find the link within the row
        link = row.find('a')
        if not link:
            continue

        link_text = link.text.strip()
        # Filter for folder links we are interested in
        if link_text.startswith('TSGR_') and not link.get('class'):
            href = link.get('href')
            if not href:
                continue

            # The date is in the next 'td' element after the link's parent 'td'
            date_td = link.find_parent('td').find_next_sibling('td')
            if not date_td:
                continue

            date_str = date_td.text.strip()
            try:
                # Parse date format like "2025/08/10 22:07"
                mod_date = datetime.strptime(date_str, "%Y/%m/%d %H:%M")
                folders_with_dates.append((mod_date, href))
            except ValueError:
                print(f"Could not parse date '{date_str}' for folder {link_text}")
                continue

    # Sort the list of tuples by date (the first element), descending
    folders_with_dates.sort(key=lambda x: x[0], reverse=True)

    # Extract just the URLs from the sorted list
    sorted_links = [href for mod_date, href in folders_with_dates]
    
    print(f"Found and sorted {len(sorted_links)} meeting folders.")
    return sorted_links

def find_excel_in_docs(docs_url):
    """
    Finds and downloads the main .xlsx file from a meeting's Docs folder.
    """
    print(f"Searching for Excel file in: {docs_url}")
    try:
        response = requests.get(docs_url)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Could not access {docs_url}. Error: {e}")
        return None

    soup = BeautifulSoup(response.text, 'html.parser')
    excel_link = None
    
    for link in soup.find_all('a'):
        href = link.get('href')
        if href and href.endswith('.xlsx'):
            excel_link = href
            break

    if not excel_link:
        return None

    try:
        # The link should be a full URL, but we use urljoin for safety
        excel_full_url = urljoin(docs_url, excel_link)
        print(f"Found Excel file: {excel_full_url}")
        
        # Get the filename from the URL
        file_name = os.path.basename(excel_full_url.split('?')[0])
        local_path = os.path.join(TEMP_DIR, file_name)

        print(f"Downloading to: {local_path}")
        with requests.get(excel_full_url, stream=True) as r:
            r.raise_for_status()
            with open(local_path, 'wb') as f:
                for chunk in r.iter_content(chunk_size=8192):
                    f.write(chunk)
        
        return local_path
    except requests.exceptions.RequestException as e:
        print(f"Failed to download {excel_link}. Error: {e}")
        return None

def filter_approved_crs(excel_path, spec_number):
    """
    Filters the downloaded Excel file for approved CRs for the specified spec.
    Handles case, whitespace, and formatting issues.
    """
    print(f"Filtering CRs in: {excel_path}")
    try:
        xls = pd.ExcelFile(excel_path)
        if 'CR_Packs_List' not in xls.sheet_names:
            print(f"Error: Sheet 'CR_Packs_List' not found in {excel_path}.")
            return []

        # Load the specific sheet
        df = pd.read_excel(xls, sheet_name='CR_Packs_List')

        # Define column names based on your screenshot
        col_rp = 'CR Pack TDoc'     # Column A
        col_r4 = 'WG Tdoc'          # Column B
        col_status = 'CR Individual TSG decision'  # Column D
        col_spec = 'Spec'           # Column E

        # Verify that all required columns exist
        required_cols = [col_rp, col_r4, col_status, col_spec]
        if not all(col in df.columns for col in required_cols):
            print(f"Error: The sheet in {excel_path} is missing one or more required columns.")
            return []

        # Normalize status and spec columns for safe comparison
        status_clean = df[col_status].astype(str).str.strip().str.lower()
        spec_clean = df[col_spec].astype(str).str.strip()

        # Apply filters
        approved_filter = status_clean == 'approved'  # lowercase
        spec_filter = spec_clean == spec_number      # exact match

        filtered_df = df[approved_filter & spec_filter]

        if filtered_df.empty:
            print("No rows matched the filter criteria.")
            return []

        # Extract data from the filtered rows
        results = []
        for index, row in filtered_df.iterrows():
            rp_number = row[col_rp]
            r4_docs = row[col_r4]
            if isinstance(r4_docs, str):
                for r4_doc in r4_docs.replace(' ', '').split(','):
                    if r4_doc.strip():  # avoid empty strings
                        results.append((str(rp_number), r4_doc.strip()))

        print(f"Found {len(results)} relevant CR(s) in {excel_path}")
        
        # Process all CRs (no limit)
        return results

    except Exception as e:
        print(f"An unexpected error occurred while processing {excel_path}: {e}")
        import traceback
        traceback.print_exc()
        return []

def process_rp_archive(docs_url, rp_number, r4_doc_name, clauses_db):
    """
    Downloads the RP archive, extracts the R4 doc, and triggers the search.
    """
    if not rp_number or not isinstance(rp_number, str):
        print(f"Invalid rp_number: {rp_number}")
        return None

    zip_url = urljoin(docs_url, rp_number + '.zip')
    zip_local_path = os.path.join(TEMP_DIR, rp_number + '.zip')
    extracted_docx_path = None
    result = None  # Changed from is_relevant to result to handle tuple

    try:
        # Download the zip file with a longer timeout
        print(f"Downloading archive: {zip_url}")
        response = requests.get(zip_url, stream=True, timeout=60)  # 60 second timeout
        response.raise_for_status()
        
        with open(zip_local_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:  # Filter out keep-alive chunks
                    f.write(chunk)

        # Verify that the downloaded file is actually a valid zip
        print(f"Verifying downloaded file: {zip_local_path}")
        if os.path.getsize(zip_local_path) == 0:
            print(f"Error: Downloaded file is empty: {zip_local_path}")
            return None
        
        # Process the downloaded zip file
        with zipfile.ZipFile(zip_local_path) as z:
            target_docx_name = r4_doc_name + '.docx'
            file_in_zip = None
            # Find a case-insensitive match for the docx file
            # Modified to search for the R4 document name anywhere in the filename
            for name in z.namelist():
                # Check if the R4 document name is contained in the filename
                # This handles cases like "38101-1_CR2917_(Rel-19)_R4-2509864_BasedOnCatFRev.docx"
                if r4_doc_name.lower() in name.lower() and name.lower().endswith('.docx'):
                    file_in_zip = name
                    break
            
            if file_in_zip:
                print(f"Found {file_in_zip} in archive. Extracting...")
                extracted_docx_path = z.extract(file_in_zip, path=TEMP_DIR)
                result = search_docx_for_clauses(extracted_docx_path, clauses_db)
            else:
                # Check if there are .zip files in the archive that might contain the .docx file
                zip_files_in_zip = [name for name in z.namelist() if name.lower().endswith('.zip')]
                docx_found = False
                
                # If we find .zip files, extract and process them
                for zip_file in zip_files_in_zip:
                    print(f"Found zip file {zip_file} in archive. Extracting...")
                    extracted_inner_zip_path = z.extract(zip_file, path=TEMP_DIR)
                    
                    # Process the inner zip file
                    try:
                        with zipfile.ZipFile(extracted_inner_zip_path) as inner_z:
                            for name in inner_z.namelist():
                                # Check if the R4 document name is contained in the filename
                                if r4_doc_name.lower() in name.lower() and name.lower().endswith('.docx'):
                                    print(f"Found {name} in inner zip file. Extracting...")
                                    extracted_docx_path = inner_z.extract(name, path=TEMP_DIR)
                                    result = search_docx_for_clauses(extracted_docx_path, clauses_db)
                                    docx_found = True
                                    break
                    except Exception as inner_e:
                        print(f"Error processing inner zip file {zip_file}: {inner_e}")
                    
                    # Clean up the inner zip file
                    if os.path.exists(extracted_inner_zip_path):
                        try:
                            os.remove(extracted_inner_zip_path)
                        except:
                            pass  # Ignore errors
                    
                    if docx_found:
                        break
                
                if not docx_found:
                    print(f"Could not find {target_docx_name} in {zip_local_path}")
                    # List available files for debugging
                    print(f"Available files in archive: {z.namelist()}")

    except requests.exceptions.Timeout:
        print(f"Timeout occurred while downloading {zip_url}")
    except requests.exceptions.RequestException as e:
        print(f"Failed to download {zip_url}. Error: {e}")
    except zipfile.BadZipFile:
        print(f"Error: {zip_local_path} is not a valid zip file.")
        # Let's check the file content
        try:
            with open(zip_local_path, 'rb') as f:
                first_bytes = f.read(100)
                print(f"First 100 bytes of file: {first_bytes}")
        except Exception as e:
            print(f"Could not read file for debugging: {e}")
    except Exception as e:
        print(f"An unexpected error occurred in process_rp_archive: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Clean up extracted and downloaded files
        if extracted_docx_path and os.path.exists(extracted_docx_path):
            try:
                os.remove(extracted_docx_path)
            except:
                pass  # Ignore errors when removing extracted file
        if os.path.exists(zip_local_path):
            try:
                os.remove(zip_local_path)
            except:
                pass  # Ignore errors when removing zip file
            
    return result  # Return the result (either tuple or None)

def search_docx_for_clauses(docx_path, clauses_db):
    """
    Searches a .docx file for the 'Clauses Affected' section and checks against the database.
    Also extracts the 'Summary of change' when a matching clause is found.
    Returns a tuple (matching_clause, summary_of_change) if found, otherwise returns None.
    """
    try:
        doc = docx.Document(docx_path)
        
        # Get all paragraphs and table content
        all_paragraphs = []
        for p in doc.paragraphs:
            all_paragraphs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.append(cell.text)

        # Debug: Print the first few paragraphs to understand document structure
        print(f"DEBUG: First few paragraphs in {docx_path}:")
        for i, para in enumerate(all_paragraphs[:10]):  # Print first 10 paragraphs
            if para.strip():
                print(f"  {i}: {para[:100]}...")  # First 100 characters of each paragraph

        # Find "Clauses Affected" section and potential matching clauses
        clauses_affected_idx = -1
        potential_clauses = []
        
        # First, find all matching clauses
        for i, text in enumerate(all_paragraphs):
            text_lower = text.lower()
            if 'clauses affected' in text_lower:
                clauses_affected_idx = i
                print(f"DEBUG: Found 'clauses affected' at paragraph {i}: {text[:100]}...")
                
                # Once found, the clause numbers could be in the same text block or the next few
                search_area = text
                # Check next few paragraphs for clauses
                for j in range(1, 6):  # Look at next 5 paragraphs to be more thorough
                    if i + j < len(all_paragraphs):
                        next_text = all_paragraphs[i + j]
                        search_area += " " + next_text
                        print(f"DEBUG: Also checking paragraph {i+j} for clauses: {next_text[:100]}...")
                
                # Use regex to find potential clause numbers (e.g., 4.1, 5.3.2, 7.1a)
                # This pattern looks for sequences of numbers and letters separated by dots.
                found_clauses = re.findall(r'[\d\w\.]+\.[\d\w]+', search_area)
                print(f"DEBUG: Found potential clauses in search area: {found_clauses}")
                
                for clause in found_clauses:
                    # Clean up the extracted clause number
                    cleaned_clause = clause.strip('., ')
                    if cleaned_clause in clauses_db:
                        print(f"Found matching clause: {cleaned_clause} in {docx_path}")
                        if cleaned_clause not in [pc[0] for pc in potential_clauses]:  # Avoid duplicates
                            potential_clauses.append((cleaned_clause, i))  # Store with index for context

        # If we found matching clauses, now look for the summary
        if potential_clauses:
            # Find the summary of change in the document
            summary_of_change = ""
            
            # First, check if we can find a "Summary of change" section
            summary_found = False
            for i, text in enumerate(all_paragraphs):
                text_lower = text.lower()
                if 'summary of change' in text_lower or 'summary of the change' in text_lower:
                    print(f"DEBUG: Found 'summary of change' at paragraph {i}: {text}")
                    summary_found = True
                    # Found the summary section, now extract the content that follows
                    start_idx = i + 1
                    summary_lines = []
                    # Look for the next few paragraphs after the header
                    print(f"DEBUG: Looking for content after paragraph {i}, starting at paragraph {start_idx}")
                    for j in range(start_idx, min(len(all_paragraphs), start_idx + 10)):  # Check next 10 paragraphs
                        para_text = all_paragraphs[j].strip()
                        print(f"DEBUG: Examining paragraph {j}: '{para_text}'")
                        
                        # Check if this is just the header line repeated
                        if para_text.lower() in ['summary of change:', 'summary of the change:', 'summary of change', 'summary of the change']:
                            print(f"DEBUG: Skipping repeated header: {para_text}")
                            continue  # Skip if it's the header line repeated
                        
                        # If we encounter a blank line or whitespace, continue but note it
                        if not para_text or para_text.isspace():
                            print(f"DEBUG: Encountered blank line at paragraph {j}")
                            continue
                        
                        # Check if this looks like a new section header (all caps, or ending with colon that isn't part of content)
                        is_section_header = (para_text.isupper() and len(para_text) < 100 and para_text.lower() not in ['summary of change', 'summary of the change', 'description of change', 'details of change', 'explanation of change'])
                        
                        # Check if it ends with a colon but is not a continuation of summary content
                        # Exclude headers or section indicators, but allow text that naturally ends with colons
                        ends_with_colon = (para_text.endswith(':') and len(para_text) < 50 and 
                                         para_text.lower() not in ['summary of change', 'summary of the change', 'description of change', 'table of changes', 'list of changes', 'overview', 'section'] and
                                         not any(keyword in para_text.lower() for keyword in ['title', 'heading', 'section', 'chapter', 'clause', 'item', 'specification', 'requirement']))
                        
                        if is_section_header or ends_with_colon:
                            # This is likely a new section header, so stop
                            print(f"DEBUG: Stopping summary extraction at potential section header: {para_text}")
                            break
                        
                        # Add this paragraph to our summary if it's not a duplicate of the last one added
                        if not summary_lines or summary_lines[-1] != para_text:
                            summary_lines.append(para_text)
                            print(f"DEBUG: Added to summary: {para_text[:50]}...")
                        else:
                            print(f"DEBUG: Skipping duplicate content: {para_text[:50]}...")
                    
                    summary_of_change = "\n".join(summary_lines).strip()
                    print(f"DEBUG: Final extracted summary of change: '{summary_of_change}'")
                    break
            
            # If no summary was found with the header, try to get content near the clauses affected
            if not summary_found and clauses_affected_idx != -1:
                print(f"DEBUG: No explicit 'summary of change' found, looking near clauses affected at index {clauses_affected_idx}")
                # Look for content in the paragraphs immediately following the clauses affected
                for i in range(clauses_affected_idx + 1, min(len(all_paragraphs), clauses_affected_idx + 15)):
                    para_text = all_paragraphs[i].strip()
                    if para_text and 'clauses affected' not in para_text.lower():
                        # Check if this looks like a new section header
                        if not ((para_text.isupper() and len(para_text) < 100) or para_text.endswith(':')):
                            if len(summary_of_change) > 0:
                                summary_of_change += "\n"
                            summary_of_change += para_text
                        else:
                            # This looks like a new section header, so stop
                            print(f"DEBUG: Stopping content extraction after clauses affected at potential header: {para_text}")
                            break
            
            print(f"DEBUG: Final summary of change for matching clause: '{summary_of_change[:100]}...'")  # First 100 chars
            # Return the first matching clause and its associated summary
            return (potential_clauses[0][0], summary_of_change)
        
        return None  # Return None if no match found

    except Exception as e:
        print(f"Error reading docx file {docx_path}: {e}")
        import traceback
        traceback.print_exc()
        return None

def single_folder_test(folder_url):
    """
    Runs a focused, end-to-end test on a single folder and the first valid CR pair.
    """
    print(f"--- Starting Focused Test on: {folder_url} ---")
    
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)

    # 1. Find and download the Excel file
    docs_url = urljoin(folder_url + '/', 'Docs/')
    excel_file_path = find_excel_in_docs(docs_url)

    if not excel_file_path:
        print("Test Failed: Could not find or download the Excel file.")
        return

    # 2. Filter for relevant CRs
    relevant_crs = filter_approved_crs(excel_file_path, SPEC_NUMBER)

    if not relevant_crs:
        print("Test Failed: No relevant CRs found in the Excel file.")
        return

    # 3. Process only the FIRST relevant CR pair
    rp_number, r4_doc_name = relevant_crs[0]
    print(f"\n--- Testing LIVE archive processing for first pair ---")
    print(f"[Progress 1/1] Processing RP: {rp_number}, R4: {r4_doc_name}")
    
    is_relevant = process_rp_archive(docs_url, rp_number, r4_doc_name, CLAUSES_DATABASE)
    
    if is_relevant:
        print(f"\nSuccess! Found a matching clause in {r4_doc_name} for RP {rp_number}.")
    else:
        print(f"\nTest complete. No matching clauses found in the first processed document.")




def run_complete_workflow():
    """
    Runs the complete workflow across all meetings and generates the final output.
    """
    print(f"[{datetime.now().strftime('%H:%M:%S')}] Starting complete 3GPP automation workflow...")
    
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)

    # 1. Get all meeting folders sorted by date (most recent first)
    meeting_folders = get_sorted_meeting_folders(BASE_URL)
    
    if not meeting_folders:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] No meeting folders found. Exiting.")
        return

    print(f"[{datetime.now().strftime('%H:%M:%S')}] Found {len(meeting_folders)} meeting folders to process.")
    
    # List to store all matching results
    all_matches = []
    
    # 2. Process each meeting folder
    for i, folder_href in enumerate(meeting_folders):
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Processing folder {i+1}/{len(meeting_folders)}: {folder_href}")
        
        folder_url = urljoin(BASE_URL + '/', folder_href)
        docs_url = urljoin(folder_url + '/', 'Docs/')
        
        # Find and download the Excel file
        excel_file_path = find_excel_in_docs(docs_url)
        
        if not excel_file_path:
            print(f"[{datetime.now().strftime('%H:%M:%S')}] Could not find Excel file in {docs_url}, skipping...")
            continue

        # Filter for relevant CRs
        relevant_crs = filter_approved_crs(excel_file_path, SPEC_NUMBER)

        if not relevant_crs:
            print(f"[{datetime.now().strftime('%H:%M:%S')}] No relevant CRs found, skipping...")
            continue

        print(f"[{datetime.now().strftime('%H:%M:%S')}] Found {len(relevant_crs)} relevant CRs to process...")
        
        # Process each relevant CR in this folder
        for j, (rp_number, r4_doc_name) in enumerate(relevant_crs):
            print(f"[{datetime.now().strftime('%H:%M:%S')}]   Processing {j+1}/{len(relevant_crs)} - RP: {rp_number}, R4: {r4_doc_name}")
            
            result = process_rp_archive(docs_url, rp_number, r4_doc_name, CLAUSES_DATABASE)
            
            if result:  # If a result was returned (not None)
                matching_clause, summary_of_change = result
                print(f"[{datetime.now().strftime('%H:%M:%S')}]     -> Match found! Clause: {matching_clause}")
                # Add match information to results with the actual matching clause and summary
                all_matches.append({
                    'Meeting Folder': folder_href,
                    'RP Number': rp_number,
                    'R4 Document': r4_doc_name,
                    'Matching Clause': matching_clause,
                    'Summary of Change': summary_of_change
                })
            else:
                print(f"[{datetime.now().strftime('%H:%M:%S')}]     -> No matching clauses found.")
        
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Completed processing {folder_href}")
    
    # 3. Generate final output file
    if all_matches:
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] Found {len(all_matches)} total matches. Generating output file: {OUTPUT_FILE}")
        
        # Create a DataFrame with the results
        df = pd.DataFrame(all_matches)
        
        # Save to Excel file
        with pd.ExcelWriter(OUTPUT_FILE, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Matches')
        
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Results saved to {OUTPUT_FILE}")
    else:
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] No matches found across all meetings.")
        # Create an empty sheet anyway
        df = pd.DataFrame(columns=['Meeting Folder', 'RP Number', 'R4 Document', 'Matching Clause', 'Summary of Change'])
        with pd.ExcelWriter(OUTPUT_FILE, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Matches')
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Empty results file created: {OUTPUT_FILE}")


if __name__ == "__main__":
    # Find and process only the first meeting folder (after sorting by date) that has an Excel file in its Docs folder
    print(f"[{datetime.now().strftime('%H:%M:%S')}] Starting latest meeting folder processing...")
    
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)

    # Get all meeting folders sorted by date (most recent first)
    meeting_folders = get_sorted_meeting_folders(BASE_URL)
    
    if not meeting_folders:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] No meeting folders found. Exiting.")
        exit(1)

    print(f"[{datetime.now().strftime('%H:%M:%S')}] Found {len(meeting_folders)} meeting folders, searching for the first one with an Excel file...")
    
    # Find the first meeting folder that has an Excel file in its Docs folder
    target_folder = None
    excel_file_path = None
    for folder_href in meeting_folders:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Checking folder: {folder_href}")
        
        folder_url = urljoin(BASE_URL + '/', folder_href)
        docs_url = urljoin(folder_url + '/', 'Docs/')
        
        # Find and download the Excel file
        excel_path = find_excel_in_docs(docs_url)
        
        if excel_path:
            target_folder = folder_href
            excel_file_path = excel_path
            print(f"[{datetime.now().strftime('%H:%M:%S')}] Found Excel file in: {target_folder}")
            break
        else:
            print(f"[{datetime.now().strftime('%H:%M:%S')}] No Excel file found in {folder_href}, checking next...")
    
    if not target_folder:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] No meeting folder with Excel file found. Exiting.")
        exit(1)

    # List to store all matching results
    all_matches = []
    
    print(f"[{datetime.now().strftime('%H:%M:%S')}] Processing target folder: {target_folder}")
    
    folder_url = urljoin(BASE_URL + '/', target_folder)
    docs_url = urljoin(folder_url + '/', 'Docs/')
    
    # Filter for relevant CRs using the excel file that was already downloaded
    relevant_crs = filter_approved_crs(excel_file_path, SPEC_NUMBER)

    if not relevant_crs:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] No relevant CRs found in the latest meeting with Excel file.")
    else:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Found {len(relevant_crs)} relevant CRs to process in the latest meeting...")
        
        # Process each relevant CR in this folder
        for j, (rp_number, r4_doc_name) in enumerate(relevant_crs):
            print(f"[{datetime.now().strftime('%H:%M:%S')}]   Processing {j+1}/{len(relevant_crs)} - RP: {rp_number}, R4: {r4_doc_name}")
            
            result = process_rp_archive(docs_url, rp_number, r4_doc_name, CLAUSES_DATABASE)
            
            if result:  # If a result was returned (not None)
                matching_clause, summary_of_change = result
                print(f"[{datetime.now().strftime('%H:%M:%S')}]     -> Match found! Clause: {matching_clause}")
                # Add match information to results with the actual matching clause and summary
                all_matches.append({
                    'Meeting Folder': target_folder,
                    'RP Number': rp_number,
                    'R4 Document': r4_doc_name,
                    'Matching Clause': matching_clause,
                    'Summary of Change': summary_of_change
                })
            else:
                print(f"[{datetime.now().strftime('%H:%M:%S')}]     -> No matching clauses found.")
            
            # Show progress: completed and pending
            completed = j + 1
            pending = len(relevant_crs) - completed
            print(f"[{datetime.now().strftime('%H:%M:%S')}]   Progress: {completed} completed, {pending} pending")
        
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Completed processing latest meeting with Excel file: {target_folder}")

    # Generate final output file
    if all_matches:
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] Found {len(all_matches)} total matches. Generating output file: {OUTPUT_FILE}")
        
        # Create a DataFrame with the results
        df = pd.DataFrame(all_matches)
        
        # Save to Excel file
        with pd.ExcelWriter(OUTPUT_FILE, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Matches')
        
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Results saved to {OUTPUT_FILE}")
    else:
        print(f"\n[{datetime.now().strftime('%H:%M:%S')}] No matches found.")
        # Create an empty sheet anyway
        df = pd.DataFrame(columns=['Meeting Folder', 'RP Number', 'R4 Document', 'Matching Clause', 'Summary of Change'])
        with pd.ExcelWriter(OUTPUT_FILE, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Matches')
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Empty results file created: {OUTPUT_FILE}")
